"""Document parser for multiple formats (DDL, Word, HTML, PDF, Excel)."""
import os
import re
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from abc import ABC, abstractmethod
import sqlparse
from src.utils.logger import get_logger
from src.utils.helpers import get_file_extension

logger = get_logger(__name__)


@dataclass
class ParsedContent:
    """Parsed document content."""

    source_type: str  # 'ddl', 'html', 'excel', 'word', 'pdf'
    tables: Dict[str, List[Dict[str, Any]]]  # table_name -> columns
    raw_text: str
    metadata: Dict[str, Any]


class DocumentParserBase(ABC):
    """Base class for document parsers."""

    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        """Check if parser can handle this file."""
        pass

    @abstractmethod
    def parse(self, file_path: str) -> ParsedContent:
        """Parse document and return content."""
        pass


class DDLParser(DocumentParserBase):
    """Parser for SQL DDL statements."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is SQL DDL."""
        ext = get_file_extension(file_path)
        return ext in [".sql", ".ddl"]

    def parse(self, file_path: str) -> ParsedContent:
        """Parse SQL DDL file."""
        logger.info(f"Parsing DDL file: {file_path}")

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        tables = self._extract_tables(content)

        return ParsedContent(
            source_type="ddl",
            tables=tables,
            raw_text=content,
            metadata={"file": file_path, "format": "SQL DDL"},
        )

    def _extract_tables(self, sql_content: str) -> Dict[str, List[Dict[str, Any]]]:
        """Extract table definitions from SQL."""
        tables = {}

        # Split by CREATE TABLE statements
        create_table_pattern = r"CREATE\s+TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?(\w+)\s*\((.*?)\);"
        matches = re.finditer(create_table_pattern, sql_content, re.IGNORECASE | re.DOTALL)

        for match in matches:
            table_name = match.group(1)
            columns_text = match.group(2)
            columns = self._parse_columns(columns_text)
            tables[table_name] = columns
            logger.debug(f"Extracted table: {table_name} with {len(columns)} columns")

        return tables

    def _parse_columns(self, columns_text: str) -> List[Dict[str, Any]]:
        """Parse column definitions."""
        columns = []

        # Split by comma, but be careful with nested structures
        column_defs = self._split_columns(columns_text)

        for col_def in column_defs:
            col_def = col_def.strip()
            if not col_def or col_def.startswith("PRIMARY") or col_def.startswith("FOREIGN"):
                continue

            parts = col_def.split(None, 2)
            if len(parts) < 2:
                continue

            col_name = parts[0]
            col_type = parts[1]

            # Extract constraints
            nullable = "NOT NULL" not in col_def.upper()
            is_primary = "PRIMARY KEY" in col_def.upper()
            is_unique = "UNIQUE" in col_def.upper()
            has_default = "DEFAULT" in col_def.upper()

            columns.append(
                {
                    "name": col_name,
                    "type": col_type,
                    "nullable": nullable,
                    "primary_key": is_primary,
                    "unique": is_unique,
                    "has_default": has_default,
                    "raw": col_def,
                }
            )

        return columns

    def _split_columns(self, columns_text: str) -> List[str]:
        """Split column definitions by comma, respecting nested structures."""
        columns = []
        current = ""
        paren_count = 0

        for char in columns_text:
            if char == "(":
                paren_count += 1
                current += char
            elif char == ")":
                paren_count -= 1
                current += char
            elif char == "," and paren_count == 0:
                columns.append(current)
                current = ""
            else:
                current += char

        if current:
            columns.append(current)

        return columns


class ExcelParser(DocumentParserBase):
    """Parser for Excel files."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is Excel."""
        ext = get_file_extension(file_path)
        return ext in [".xlsx", ".xls"]

    def parse(self, file_path: str) -> ParsedContent:
        """Parse Excel file."""
        logger.info(f"Parsing Excel file: {file_path}")

        try:
            import openpyxl
            import pandas as pd
        except ImportError:
            raise ImportError("openpyxl and pandas required for Excel parsing")

        workbook = openpyxl.load_workbook(file_path)
        tables = {}

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            columns = self._extract_columns_from_sheet(sheet)
            if columns:
                tables[sheet_name] = columns

        return ParsedContent(
            source_type="excel",
            tables=tables,
            raw_text="",
            metadata={"file": file_path, "format": "Excel", "sheets": list(workbook.sheetnames)},
        )

    def _extract_columns_from_sheet(self, sheet) -> List[Dict[str, Any]]:
        """Extract column info from Excel sheet."""
        columns = []

        # Assume first row contains headers
        first_row = list(sheet.iter_rows(min_row=1, max_row=1, values_only=False))
        if not first_row:
            return columns

        for cell in first_row[0]:
            if cell.value:
                columns.append(
                    {
                        "name": str(cell.value),
                        "type": "VARCHAR",  # Default type for Excel
                        "nullable": True,
                        "primary_key": False,
                        "unique": False,
                        "has_default": False,
                        "raw": str(cell.value),
                    }
                )

        return columns


class HTMLParser(DocumentParserBase):
    """Parser for HTML files."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is HTML."""
        ext = get_file_extension(file_path)
        return ext in [".html", ".htm"]

    def parse(self, file_path: str) -> ParsedContent:
        """Parse HTML file."""
        logger.info(f"Parsing HTML file: {file_path}")

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 required for HTML parsing")

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")

        tables = {}
        table_elements = soup.find_all("table")

        for idx, table in enumerate(table_elements):
            table_name = f"table_{idx}"
            columns = self._extract_columns_from_table(table)
            if columns:
                tables[table_name] = columns

        return ParsedContent(
            source_type="html",
            tables=tables,
            raw_text=str(soup),
            metadata={"file": file_path, "format": "HTML", "table_count": len(table_elements)},
        )

    def _extract_columns_from_table(self, table) -> List[Dict[str, Any]]:
        """Extract column info from HTML table."""
        columns = []

        # Look for header row
        header_row = table.find("tr")
        if not header_row:
            return columns

        # Check for <th> first, then <td>
        headers = header_row.find_all("th")
        if not headers:
            headers = header_row.find_all("td")

        for header in headers:
            col_name = header.get_text().strip()
            if col_name:
                columns.append(
                    {
                        "name": col_name,
                        "type": "VARCHAR",
                        "nullable": True,
                        "primary_key": False,
                        "unique": False,
                        "has_default": False,
                        "raw": col_name,
                    }
                )

        return columns


class WordParser(DocumentParserBase):
    """Parser for Word documents."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is Word document."""
        ext = get_file_extension(file_path)
        return ext in [".docx", ".doc"]

    def parse(self, file_path: str) -> ParsedContent:
        """Parse Word document."""
        logger.info(f"Parsing Word file: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for Word parsing")

        doc = Document(file_path)
        tables = {}
        raw_text = []

        # Extract text
        for para in doc.paragraphs:
            if para.text.strip():
                raw_text.append(para.text)

        # Extract tables
        for idx, table in enumerate(doc.tables):
            table_name = f"table_{idx}"
            columns = self._extract_columns_from_table(table)
            if columns:
                tables[table_name] = columns

        return ParsedContent(
            source_type="word",
            tables=tables,
            raw_text="\n".join(raw_text),
            metadata={"file": file_path, "format": "Word", "table_count": len(doc.tables)},
        )

    def _extract_columns_from_table(self, table) -> List[Dict[str, Any]]:
        """Extract column info from Word table."""
        columns = []

        if not table.rows:
            return columns

        # First row is headers
        first_row = table.rows[0]

        for cell in first_row.cells:
            col_name = cell.text.strip()
            if col_name:
                columns.append(
                    {
                        "name": col_name,
                        "type": "VARCHAR",
                        "nullable": True,
                        "primary_key": False,
                        "unique": False,
                        "has_default": False,
                        "raw": col_name,
                    }
                )

        return columns


class PDFParser(DocumentParserBase):
    """Parser for PDF files."""

    def can_parse(self, file_path: str) -> bool:
        """Check if file is PDF."""
        ext = get_file_extension(file_path)
        return ext == ".pdf"

    def parse(self, file_path: str) -> ParsedContent:
        """Parse PDF file."""
        logger.info(f"Parsing PDF file: {file_path}")

        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber required for PDF parsing")

        tables = {}
        raw_text = []

        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                # Extract text
                text = page.extract_text()
                if text:
                    raw_text.append(text)

                # Extract tables
                extracted_tables = page.extract_tables()
                if extracted_tables:
                    for table_idx, table_data in enumerate(extracted_tables):
                        table_name = f"table_page{page_idx}_idx{table_idx}"
                        columns = self._extract_columns_from_data(table_data)
                        if columns:
                            tables[table_name] = columns

        return ParsedContent(
            source_type="pdf",
            tables=tables,
            raw_text="\n".join(raw_text),
            metadata={"file": file_path, "format": "PDF", "table_count": len(tables)},
        )

    def _extract_columns_from_data(self, table_data: List[List]) -> List[Dict[str, Any]]:
        """Extract column info from PDF table data."""
        columns = []

        if not table_data or not table_data[0]:
            return columns

        # First row is headers
        for col_name in table_data[0]:
            if col_name:
                columns.append(
                    {
                        "name": str(col_name),
                        "type": "VARCHAR",
                        "nullable": True,
                        "primary_key": False,
                        "unique": False,
                        "has_default": False,
                        "raw": str(col_name),
                    }
                )

        return columns


class DocumentParser:
    """Main document parser that routes to appropriate parser."""

    def __init__(self):
        """Initialize with all parsers."""
        self.parsers = [
            DDLParser(),
            ExcelParser(),
            HTMLParser(),
            WordParser(),
            PDFParser(),
        ]

    def parse(self, file_path: str) -> ParsedContent:
        """Parse document using appropriate parser."""
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser.parse(file_path)

        raise ValueError(f"No parser found for file: {file_path}")

    def parse_multiple(self, file_paths: List[str]) -> List[ParsedContent]:
        """Parse multiple documents."""
        return [self.parse(file_path) for file_path in file_paths]
